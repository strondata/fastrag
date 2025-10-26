"""Testes para multi-format loaders."""

import pytest
import tempfile
import shutil
from pathlib import Path

from src.rag_chatbot.components.loaders import UniversalLoader, FolderLoader
from src.rag_chatbot.interfaces import Documento


class TestUniversalLoader:
    """Testes para UniversalLoader."""
    
    @pytest.fixture
    def temp_data_dir(self):
        """Cria um diretório temporário com arquivos de teste."""
        temp_dir = tempfile.mkdtemp()
        
        # Criar arquivos de texto
        (Path(temp_dir) / "test1.txt").write_text("Conteúdo do arquivo de texto 1", encoding='utf-8')
        (Path(temp_dir) / "test2.md").write_text("# Markdown\nConteúdo markdown", encoding='utf-8')
        
        yield temp_dir
        
        # Limpar
        shutil.rmtree(temp_dir)
    
    def test_load_text_files(self, temp_data_dir):
        """Testa carregamento de arquivos de texto."""
        loader = UniversalLoader()
        documents = loader.load(temp_data_dir)
        
        # Deve ter pelo menos 2 documentos (.txt e .md)
        assert len(documents) >= 2
        
        # Verificar que os tipos estão corretos
        text_docs = [d for d in documents if d.metadata.get('type') == 'text']
        assert len(text_docs) >= 2
        
        # Verificar metadados
        for doc in documents:
            assert 'source' in doc.metadata
            assert 'path' in doc.metadata
            assert 'type' in doc.metadata
    
    def test_load_pdf_file(self, temp_data_dir):
        """Testa carregamento de arquivo PDF."""
        try:
            import fitz  # PyMuPDF
            
            # Criar um PDF simples
            pdf_path = Path(temp_data_dir) / "test.pdf"
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((50, 50), "Conteúdo do PDF de teste")
            doc.save(pdf_path)
            doc.close()
            
            # Carregar com UniversalLoader
            loader = UniversalLoader()
            documents = loader.load(temp_data_dir)
            
            # Encontrar o documento PDF
            pdf_docs = [d for d in documents if d.metadata.get('type') == 'pdf']
            assert len(pdf_docs) == 1
            
            pdf_doc = pdf_docs[0]
            assert 'Conteúdo do PDF de teste' in pdf_doc.content
            assert pdf_doc.metadata['source'] == 'test.pdf'
            assert pdf_doc.metadata['type'] == 'pdf'
            assert 'pages' in pdf_doc.metadata
            
        except ImportError:
            pytest.skip("PyMuPDF não está instalado")
    
    def test_load_docx_file(self, temp_data_dir):
        """Testa carregamento de arquivo DOCX."""
        try:
            from docx import Document as DocxDocument
            
            # Criar um DOCX simples
            docx_path = Path(temp_data_dir) / "test.docx"
            doc = DocxDocument()
            doc.add_paragraph("Parágrafo 1 do DOCX")
            doc.add_paragraph("Parágrafo 2 do DOCX")
            doc.save(docx_path)
            
            # Carregar com UniversalLoader
            loader = UniversalLoader()
            documents = loader.load(temp_data_dir)
            
            # Encontrar o documento DOCX
            docx_docs = [d for d in documents if d.metadata.get('type') == 'docx']
            assert len(docx_docs) == 1
            
            docx_doc = docx_docs[0]
            assert 'Parágrafo 1 do DOCX' in docx_doc.content
            assert 'Parágrafo 2 do DOCX' in docx_doc.content
            assert docx_doc.metadata['source'] == 'test.docx'
            assert docx_doc.metadata['type'] == 'docx'
            assert 'paragraphs' in docx_doc.metadata
            
        except ImportError:
            pytest.skip("python-docx não está instalado")
    
    def test_folder_loader_alias(self, temp_data_dir):
        """Testa que FolderLoader funciona como alias para UniversalLoader."""
        loader = FolderLoader()
        documents = loader.load(temp_data_dir)
        
        # Deve funcionar da mesma forma que UniversalLoader
        assert len(documents) >= 2
        assert all('source' in doc.metadata for doc in documents)
    
    def test_mixed_formats(self, temp_data_dir):
        """Testa carregamento de pasta com formatos mistos."""
        # temp_data_dir já tem .txt e .md
        loader = UniversalLoader()
        documents = loader.load(temp_data_dir)
        
        # Verificar que encontrou múltiplos formatos
        types = set(doc.metadata.get('type') for doc in documents)
        assert 'text' in types
        
        # Todos devem ter os metadados básicos
        for doc in documents:
            assert doc.content is not None
            assert doc.metadata.get('source') is not None
            assert doc.metadata.get('path') is not None
            assert doc.metadata.get('type') is not None
    
    def test_error_handling(self, temp_data_dir):
        """Testa tratamento de erros em arquivos problemáticos."""
        # Criar um arquivo "PDF" corrompido (apenas texto)
        corrupt_pdf = Path(temp_data_dir) / "corrupt.pdf"
        corrupt_pdf.write_text("Este não é um PDF válido", encoding='utf-8')
        
        loader = UniversalLoader()
        
        # Não deve lançar exceção, deve continuar com outros arquivos
        try:
            documents = loader.load(temp_data_dir)
            # Deve ter carregado os arquivos válidos
            assert len(documents) >= 2  # .txt e .md devem ter sido carregados
        except Exception as e:
            pytest.fail(f"Loader não tratou erro adequadamente: {e}")
